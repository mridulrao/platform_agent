"""
Standalone Image Processing Module (Azure OpenAI Edition)

Extracts images from PDFs, DOCX, and HTML documents, generates descriptions
using Azure OpenAI vision model, and inserts those descriptions back into the
document text at the correct positions.

Dependencies:
    pip install pdfplumber pillow beautifulsoup4 pydantic aiohttp python-docx instructor openai lxml

Required env_config attributes:
    LLM_AZURE_OPENAI_API_VERSION
    LLM_AZURE_OPENAI_ENDPOINT        # base URL: https://xxx.openai.azure.com/
    LLM_AZURE_OPENAI_API_KEY
    LLM_AZURE_OPENAI_DEPLOYMENT_NAME

Usage:
    from image_processor import ImageProcessor, ImageProcessingConfig

    processor = ImageProcessor()
    text = await processor.process_pdf_with_images("document.pdf")
    text = await processor.process_docx_with_images("document.docx")
    text = await processor.process_html_with_images(html_string, base_url="https://example.com")
"""

import os
import re
import base64
import hashlib
import logging
import asyncio
import aiohttp
import io
import time
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor
from urllib.parse import urljoin
from zipfile import ZipFile

import pdfplumber
from PIL import Image
from bs4 import BeautifulSoup
from pydantic import BaseModel, Field

from env_config import config as env_config
from openai import AzureOpenAI

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    docx = None
    DOCX_AVAILABLE = False

try:
    import instructor
    INSTRUCTOR_AVAILABLE = True
except ImportError:
    instructor = None
    INSTRUCTOR_AVAILABLE = False

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════════
# Data Models
# ═══════════════════════════════════════════════════════════════════════════════

class ImageProcessingConfig(BaseModel):
    """Configuration for image processing."""

    enable_image_processing: bool = Field(
        default=True,
        description="Enable or disable image processing"
    )
    min_image_size: int = Field(
        default=32,
        description="Minimum pixel dimension for images to process"
    )
    max_concurrent_images: int = Field(
        default=10,
        description="Maximum number of images to process concurrently"
    )
    context_chars: int = Field(
        default=400,
        description="Number of surrounding characters to include for context"
    )
    filter_decorative: bool = Field(
        default=True,
        description="Filter out decorative images"
    )
    filter_headers_footers: bool = Field(
        default=True,
        description="Filter out header/footer images"
    )
    max_images_per_doc: int = Field(
        default=50,
        description="Maximum number of images to process per document"
    )
    temperature: float = Field(
        default=0.1,
        description="Temperature for vision model calls"
    )
    request_timeout: int = Field(
        default=30,
        description="Timeout in seconds for HTTP requests (image downloads)"
    )
    download_max_retries: int = Field(
        default=3,
        description="Maximum retry attempts for image downloads"
    )
    vision_max_retries: int = Field(
        default=3,
        description="Maximum retry attempts for vision model calls"
    )
    retry_base_delay_sec: float = Field(
        default=1.0,
        description="Base delay for exponential backoff retries"
    )


class ImageDescriptionResponse(BaseModel):
    """Structured response from the vision model."""

    text: str = Field(description="The generated description text")
    has_content: bool = Field(
        default=True,
        description="Whether the image has meaningful content"
    )
    is_decorative: bool = Field(
        default=False,
        description="Whether the image is purely decorative"
    )
    content_type: str = Field(
        default="comprehensive",
        description=(
            "Type of content: text_extraction | action_step | dialog_box | "
            "diagram | screenshot | result_display | instructions | unknown"
        )
    )
    is_action_step: bool = Field(
        default=False,
        description="True if the image shows an action the user needs to take"
    )
    contains_instructions: bool = Field(
        default=False,
        description="True if the image contains written instructions/steps"
    )
    extracted_text: Optional[str] = Field(
        default=None,
        description="Verbatim text extracted from the image when contains_instructions=True"
    )
    action_summary: Optional[str] = Field(
        default=None,
        description="Short action summary e.g. 'Click Accept button'"
    )
    shows_result: bool = Field(
        default=False,
        description="True if image shows a result/confirmation after an action"
    )
    ui_elements: Optional[List[str]] = Field(
        default=None,
        description="Key UI elements visible: buttons, links, checkboxes, input fields"
    )


@dataclass
class TextElement:
    """Text or image element with position information."""

    y: float
    x: float
    content: str
    element_type: str = "text"          # "text" or "image"
    alt_text: str = ""
    content_type_mime: str = "image/png"


# ═══════════════════════════════════════════════════════════════════════════════
# In-memory description cache
# ═══════════════════════════════════════════════════════════════════════════════

class ImageDescriptionCache:
    """Cache image descriptions keyed by MD5 hash of raw image bytes."""

    def __init__(self, max_size: int = 1000):
        self._cache: Dict[str, ImageDescriptionResponse] = {}
        self._max_size = max_size
        self._hits = 0
        self._misses = 0
        logger.debug("ImageDescriptionCache initialized | max_size=%d", max_size)

    def _key(self, image_data: bytes) -> str:
        return hashlib.md5(image_data).hexdigest()

    def get(self, image_data: bytes) -> Optional[ImageDescriptionResponse]:
        result = self._cache.get(self._key(image_data))
        if result:
            self._hits += 1
            logger.debug("Cache HIT  | hits=%d misses=%d", self._hits, self._misses)
        else:
            self._misses += 1
            logger.debug("Cache MISS | hits=%d misses=%d", self._hits, self._misses)
        return result

    def set(self, image_data: bytes, response: ImageDescriptionResponse) -> None:
        if len(self._cache) >= self._max_size:
            evict_count = self._max_size // 2
            logger.debug(
                "Cache full (%d entries) — evicting oldest %d entries",
                len(self._cache), evict_count,
            )
            for k in list(self._cache.keys())[:evict_count]:
                del self._cache[k]
        self._cache[self._key(image_data)] = response
        logger.debug("Cache SET  | size=%d/%d", len(self._cache), self._max_size)

    def stats(self) -> str:
        total = self._hits + self._misses
        rate = (self._hits / total * 100) if total else 0
        return (
            f"Cache stats: size={len(self._cache)}/{self._max_size} | "
            f"hits={self._hits} | misses={self._misses} | hit_rate={rate:.1f}%"
        )


# ═══════════════════════════════════════════════════════════════════════════════
# Main Processor
# ═══════════════════════════════════════════════════════════════════════════════

class ImageProcessor:
    """
    Image processing pipeline using Azure OpenAI for vision descriptions.

    Extracts images from PDFs, DOCX, and HTML, generates descriptions via the
    Azure OpenAI vision model, and inserts those descriptions at the correct
    positions in the document text.

    All Azure credentials are sourced from env_config automatically.
    """

    def __init__(self, config: Optional[ImageProcessingConfig] = None):
        logger.info("=" * 60)
        logger.info("Initializing ImageProcessor")
        logger.info("=" * 60)

        self.config = config or ImageProcessingConfig()
        logger.info(
            "Config | enable_image_processing=%s | min_image_size=%d | "
            "max_concurrent_images=%d | max_images_per_doc=%d | "
            "filter_decorative=%s | filter_headers_footers=%s | temperature=%.2f",
            self.config.enable_image_processing,
            self.config.min_image_size,
            self.config.max_concurrent_images,
            self.config.max_images_per_doc,
            self.config.filter_decorative,
            self.config.filter_headers_footers,
            self.config.temperature,
        )

        self.cache = ImageDescriptionCache()
        self.thread_pool = ThreadPoolExecutor(
            max_workers=self.config.max_concurrent_images
        )
        logger.info(
            "ThreadPoolExecutor initialized | max_workers=%d",
            self.config.max_concurrent_images,
        )

        # Build the instructor-wrapped Azure OpenAI client
        if INSTRUCTOR_AVAILABLE:
            try:
                logger.info(
                    "Connecting to Azure OpenAI | endpoint=%s | deployment=%s | api_version=%s",
                    env_config.LLM_AZURE_OPENAI_ENDPOINT,
                    env_config.LLM_AZURE_OPENAI_DEPLOYMENT_NAME,
                    env_config.LLM_AZURE_OPENAI_API_VERSION,
                )
                azure_client = AzureOpenAI(
                    api_version=env_config.LLM_AZURE_OPENAI_API_VERSION,
                    azure_endpoint=env_config.LLM_AZURE_OPENAI_ENDPOINT,
                    api_key=env_config.LLM_AZURE_OPENAI_API_KEY,
                    azure_deployment=env_config.LLM_AZURE_OPENAI_DEPLOYMENT_NAME,
                )
                self._client = instructor.from_openai(azure_client)
                logger.info(
                    "✓ Azure OpenAI client ready | deployment=%s | endpoint=%s",
                    env_config.LLM_AZURE_OPENAI_DEPLOYMENT_NAME,
                    env_config.LLM_AZURE_OPENAI_ENDPOINT,
                )
            except Exception as exc:
                logger.error("✗ Failed to initialise Azure OpenAI client: %s", exc, exc_info=True)
                self._client = None
        else:
            self._client = None
            logger.warning(
                "instructor / openai not installed — vision calls disabled. "
                "Run: pip install instructor openai"
            )

        logger.info("ImageProcessor initialisation complete.")

    # ──────────────────────────────────────────────────────────────────────────
    # Public API
    # ──────────────────────────────────────────────────────────────────────────

    async def process_pdf_with_images(self, pdf_path: str) -> str:
        """
        Extract text + images from a PDF and return a merged text string.

        Images are described by the vision model and inserted at their
        approximate vertical position within each page's text flow.
        """
        logger.info("\n%s", "=" * 60)
        logger.info("PROCESSING PDF: %s", pdf_path)
        logger.info("=" * 60)
        t_start = time.monotonic()

        if not self.config.enable_image_processing:
            logger.info("Image processing disabled — extracting text only.")
            return await self._extract_pdf_text_only(pdf_path)

        pdf = None
        try:
            logger.info("Opening PDF: %s", pdf_path)
            pdf = pdfplumber.open(pdf_path)
            total_pages = len(pdf.pages)
            logger.info("PDF opened | pages=%d", total_pages)

            # Phase 1: Discover unique images across all pages
            logger.info("Phase 1: Discovering images across all pages...")
            t_phase = time.monotonic()
            image_map = self._discover_pdf_images(pdf)
            logger.info(
                "Phase 1 complete | unique_images=%d | elapsed=%.2fs",
                len(image_map), time.monotonic() - t_phase,
            )

            # Enforce per-document image limit
            if len(image_map) > self.config.max_images_per_doc:
                logger.warning(
                    "Image count %d exceeds limit %d — truncating to %d",
                    len(image_map), self.config.max_images_per_doc, self.config.max_images_per_doc,
                )
                limited = sorted(image_map.keys())[: self.config.max_images_per_doc]
                image_map = {h: image_map[h] for h in limited}

            # Phase 2: Extract raw bytes + generate descriptions
            logger.info("Phase 2: Extracting image bytes and generating descriptions...")
            t_phase = time.monotonic()
            image_cache: Dict[str, Tuple[bytes, str]] = {}
            description_cache: Dict[str, ImageDescriptionResponse] = {}
            await self._process_pdf_images_parallel(image_map, image_cache, description_cache)
            logger.info(
                "Phase 2 complete | extracted=%d | described=%d | elapsed=%.2fs",
                len(image_cache), len(description_cache), time.monotonic() - t_phase,
            )

            # Phase 3: Build per-page content with images merged by position
            logger.info("Phase 3: Merging text and image descriptions per page...")
            t_phase = time.monotonic()
            pages_content = []
            for page_index, page in enumerate(pdf.pages):
                logger.debug("Building content for page %d/%d", page_index + 1, total_pages)
                page_text = self._build_page_content_with_images(
                    page, page_index, image_map, image_cache, description_cache
                )
                pages_content.append(page_text)
                logger.debug(
                    "Page %d/%d | content_length=%d chars",
                    page_index + 1, total_pages, len(page_text),
                )

            logger.info(
                "Phase 3 complete | pages_with_content=%d | elapsed=%.2fs",
                sum(1 for p in pages_content if p.strip()), time.monotonic() - t_phase,
            )

            pdf.close()
            pdf = None

            result = "\n\n---\n\n".join(pages_content)
            logger.info(
                "✓ PDF processing complete | total_chars=%d | total_elapsed=%.2fs",
                len(result), time.monotonic() - t_start,
            )
            logger.info(self.cache.stats())
            return result

        except Exception as exc:
            logger.error(
                "✗ Error processing PDF %s: %s | elapsed=%.2fs",
                pdf_path, exc, time.monotonic() - t_start, exc_info=True,
            )
            logger.info("Falling back to text-only extraction.")
            return await self._extract_pdf_text_only(pdf_path)
        finally:
            if pdf:
                try:
                    pdf.close()
                except Exception:
                    pass

    async def process_docx_with_images(self, docx_path: str) -> str:
        """
        Extract text + images from a DOCX and return a merged text string.
        """
        logger.info("\n%s", "=" * 60)
        logger.info("PROCESSING DOCX: %s", docx_path)
        logger.info("=" * 60)
        t_start = time.monotonic()

        if not DOCX_AVAILABLE:
            logger.warning("python-docx not installed — falling back to text-only.")
            return await self._extract_docx_text_only(docx_path)

        if not self.config.enable_image_processing:
            logger.info("Image processing disabled — extracting text only.")
            return await self._extract_docx_text_only(docx_path)

        try:
            logger.info("Extracting DOCX elements (text + images)...")
            t_phase = time.monotonic()
            loop = asyncio.get_running_loop()
            elements = await loop.run_in_executor(
                self.thread_pool, self._extract_docx_elements_sync, docx_path
            )
            logger.info(
                "Element extraction complete | total_elements=%d | elapsed=%.2fs",
                len(elements), time.monotonic() - t_phase,
            )

            if not elements:
                logger.warning("No elements extracted from DOCX — returning empty string.")
                return await self._extract_docx_text_only(docx_path)

            text_elements  = [e for e in elements if e.element_type == "text"]
            image_elements = [e for e in elements if e.element_type == "image"]

            logger.info(
                "DOCX breakdown | text_elements=%d | image_elements=%d",
                len(text_elements), len(image_elements),
            )

            if not image_elements:
                logger.info("No images found in DOCX — returning text-only content.")
                return "\n".join(e.content for e in text_elements)

            if len(image_elements) > self.config.max_images_per_doc:
                logger.warning(
                    "Image count %d exceeds limit %d — truncating.",
                    len(image_elements), self.config.max_images_per_doc,
                )
                image_elements = image_elements[: self.config.max_images_per_doc]

            logger.info("Generating descriptions for %d images...", len(image_elements))
            t_phase = time.monotonic()
            image_descriptions = await self._process_docx_images(image_elements)
            logger.info(
                "Image descriptions complete | described=%d | elapsed=%.2fs",
                len(image_descriptions), time.monotonic() - t_phase,
            )

            logger.info("Merging text elements and image descriptions...")
            result = self._merge_docx_elements(text_elements, image_descriptions)

            logger.info(
                "✓ DOCX processing complete | total_chars=%d | total_elapsed=%.2fs",
                len(result), time.monotonic() - t_start,
            )
            logger.info(self.cache.stats())
            return result

        except Exception as exc:
            logger.error(
                "✗ Error processing DOCX %s: %s | elapsed=%.2fs",
                docx_path, exc, time.monotonic() - t_start, exc_info=True,
            )
            logger.info("Falling back to text-only extraction.")
            return await self._extract_docx_text_only(docx_path)

    async def process_html_with_images(
        self,
        html_content: str,
        base_url: Optional[str] = None,
    ) -> str:
        """
        Clean HTML, describe embedded/linked images, and return plain text
        with image descriptions inserted where the <img> tags appeared.
        """
        logger.info("\n%s", "=" * 60)
        logger.info("PROCESSING HTML | base_url=%s | input_length=%d chars", base_url, len(html_content))
        logger.info("=" * 60)
        t_start = time.monotonic()

        if not self.config.enable_image_processing:
            logger.info("Image processing disabled — cleaning HTML text only.")
            return self._clean_html_text_only(html_content)

        try:
            logger.info("Parsing HTML and removing script/style tags...")
            soup = BeautifulSoup(html_content, "html.parser")
            removed = len(soup(["script", "style"]))
            for el in soup(["script", "style"]):
                el.decompose()
            logger.debug("Removed %d script/style elements.", removed)

            imgs = soup.find_all("img")
            logger.info("Found %d <img> tags in HTML.", len(imgs))

            images_to_process: List[Tuple[int, Any, str]] = [
                (idx, tag, tag.get("src", ""))
                for idx, tag in enumerate(imgs)
                if tag.get("src", "")
            ]
            logger.info(
                "%d images have a src attribute and will be processed.",
                len(images_to_process),
            )

            if len(images_to_process) > self.config.max_images_per_doc:
                logger.warning(
                    "Image count %d exceeds limit %d — truncating.",
                    len(images_to_process), self.config.max_images_per_doc,
                )
                images_to_process = images_to_process[: self.config.max_images_per_doc]

            image_descriptions: Dict[int, str] = {}
            success_count = 0
            cache_hit_count = 0
            fallback_count = 0
            error_count = 0

            for position, img_tag, src in images_to_process:
                alt_text = img_tag.get("alt", "")
                src_preview = src[:80] + "..." if len(src) > 80 else src
                logger.info(
                    "Processing image %d/%d | src=%s | alt=%r",
                    position + 1, len(images_to_process), src_preview, alt_text,
                )

                try:
                    t_img = time.monotonic()
                    image_data = await self._get_image_data(src, base_url)

                    if not image_data:
                        logger.warning(
                            "Image %d: No data retrieved — using fallback description.", position
                        )
                        fallback = self._format_fallback_description(alt_text)
                        if fallback:
                            image_descriptions[position] = fallback
                        fallback_count += 1
                        continue

                    logger.debug("Image %d: Retrieved %d bytes.", position, len(image_data))

                    cached = self.cache.get(image_data)
                    if cached:
                        logger.info("Image %d: Cache hit — skipping vision call.", position)
                        formatted = self._format_image_description(cached, alt_text)
                        if formatted:
                            image_descriptions[position] = formatted
                        cache_hit_count += 1
                        continue

                    mime_type = self._guess_mime_type(src)
                    b64 = base64.b64encode(image_data).decode()
                    data_uri = f"data:{mime_type};base64,{b64}"

                    logger.info(
                        "Image %d: Calling vision model | mime=%s | data_size=%d bytes",
                        position, mime_type, len(image_data),
                    )
                    description = await self._generate_image_description_async(data_uri)

                    if description:
                        self.cache.set(image_data, description)
                        formatted = self._format_image_description(description, alt_text)
                        if formatted:
                            image_descriptions[position] = formatted
                        success_count += 1
                        logger.info(
                            "Image %d: ✓ Described | content_type=%s | is_action=%s | "
                            "has_instructions=%s | elapsed=%.2fs",
                            position, description.content_type,
                            description.is_action_step, description.contains_instructions,
                            time.monotonic() - t_img,
                        )
                    else:
                        logger.warning("Image %d: Vision model returned no description — using fallback.", position)
                        image_descriptions[position] = self._format_fallback_description(alt_text)
                        fallback_count += 1

                except Exception as exc:
                    error_count += 1
                    logger.warning(
                        "Image %d: ✗ Error during processing: %s | using fallback.",
                        position, exc,
                    )
                    image_descriptions[position] = self._format_fallback_description(alt_text)

            logger.info(
                "Image processing summary | total=%d | success=%d | cache_hits=%d | "
                "fallbacks=%d | errors=%d",
                len(images_to_process), success_count, cache_hit_count,
                fallback_count, error_count,
            )

            logger.info("Building final plain-text output from HTML tree...")
            result = self._clean_html_with_descriptions(soup, image_descriptions)

            logger.info(
                "✓ HTML processing complete | output_length=%d chars | total_elapsed=%.2fs",
                len(result), time.monotonic() - t_start,
            )
            logger.info(self.cache.stats())
            return result

        except Exception as exc:
            logger.error(
                "✗ Error processing HTML: %s | elapsed=%.2fs",
                exc, time.monotonic() - t_start, exc_info=True,
            )
            logger.info("Falling back to text-only HTML cleaning.")
            return self._clean_html_text_only(html_content)

    def shutdown(self) -> None:
        """Release thread pool resources."""
        logger.info("Shutting down ImageProcessor thread pool...")
        self.thread_pool.shutdown(wait=True)
        logger.info(self.cache.stats())
        logger.info("✓ ImageProcessor shutdown complete.")

    # ──────────────────────────────────────────────────────────────────────────
    # PDF internals
    # ──────────────────────────────────────────────────────────────────────────

    def _discover_pdf_images(
        self, pdf: pdfplumber.PDF
    ) -> Dict[str, List[Tuple[int, Tuple[float, float, float, float], dict]]]:
        """
        Scan every page and collect unique images keyed by MD5 of stream data.
        """
        image_map: Dict[
            str, List[Tuple[int, Tuple[float, float, float, float], dict]]
        ] = {}

        total_raw = 0
        skipped_too_small = 0
        skipped_no_stream = 0

        for page_index, page in enumerate(pdf.pages):
            page_imgs = page.images or []
            logger.debug("Page %d: found %d raw image entries.", page_index + 1, len(page_imgs))
            total_raw += len(page_imgs)

            for img in page_imgs:
                x0, y0, x1, y1 = img["x0"], img["top"], img["x1"], img["bottom"]
                w, h = x1 - x0, y1 - y0

                if w < self.config.min_image_size and h < self.config.min_image_size:
                    logger.debug(
                        "Page %d: Skipping tiny image (%.1f x %.1f px).", page_index + 1, w, h
                    )
                    skipped_too_small += 1
                    continue

                stream = img.get("stream")
                if stream is None:
                    logger.debug("Page %d: Skipping image with no stream.", page_index + 1)
                    skipped_no_stream += 1
                    continue

                try:
                    raw = stream.get_data()
                    img_hash = hashlib.md5(raw).hexdigest()
                except Exception as exc:
                    img_hash = f"p{page_index}_id{id(img)}"
                    logger.debug(
                        "Page %d: Could not hash image stream (%s) — using fallback key.",
                        page_index + 1, exc,
                    )

                image_map.setdefault(img_hash, []).append(
                    (page_index, (x0, y0, x1, y1), img)
                )

        logger.info(
            "Image discovery | total_raw=%d | skipped_tiny=%d | skipped_no_stream=%d | unique=%d",
            total_raw, skipped_too_small, skipped_no_stream, len(image_map),
        )
        return image_map

    async def _process_pdf_images_parallel(
        self,
        image_map: Dict[str, List[Tuple[int, Tuple[float, float, float, float], dict]]],
        image_cache: Dict[str, Tuple[bytes, str]],
        description_cache: Dict[str, ImageDescriptionResponse],
    ) -> None:
        """Extract bytes and generate descriptions in parallel batches."""
        unique_hashes = list(image_map.keys())
        if not unique_hashes:
            logger.info("No images to process — skipping parallel processing.")
            return

        batch_size = self.config.max_concurrent_images
        total_batches = (len(unique_hashes) + batch_size - 1) // batch_size
        logger.info(
            "Starting parallel image processing | total=%d | batch_size=%d | batches=%d",
            len(unique_hashes), batch_size, total_batches,
        )

        for batch_start in range(0, len(unique_hashes), batch_size):
            batch = unique_hashes[batch_start : batch_start + batch_size]
            batch_num = batch_start // batch_size + 1
            logger.info(
                "Batch %d/%d | images=%d | extracting bytes...",
                batch_num, total_batches, len(batch),
            )
            t_batch = time.monotonic()
            loop = asyncio.get_running_loop()

            # ── Parallel extraction ───────────────────────────────────────────
            extract_pairs = []
            for img_hash in batch:
                _, _, img_dict = image_map[img_hash][0]
                task = loop.run_in_executor(
                    self.thread_pool, self._extract_image_from_pdf, img_dict
                )
                extract_pairs.append((img_hash, task))

            extracted_count = 0
            for img_hash, task in extract_pairs:
                try:
                    img_bytes, img_ext = await task
                    if img_bytes:
                        image_cache[img_hash] = (img_bytes, img_ext)
                        extracted_count += 1
                        logger.debug(
                            "Extracted %s | format=%s | size=%d bytes",
                            img_hash[:8], img_ext, len(img_bytes),
                        )
                    else:
                        logger.warning("Extraction returned no bytes for hash %s.", img_hash[:8])
                except Exception as exc:
                    logger.error(
                        "✗ Extraction failed for hash %s: %s", img_hash[:8], exc
                    )

            logger.info(
                "Batch %d/%d | extraction done | extracted=%d/%d",
                batch_num, total_batches, extracted_count, len(batch),
            )

            # ── Parallel description generation ───────────────────────────────
            desc_tasks = []
            cache_hits = 0
            for img_hash in batch:
                if img_hash not in image_cache:
                    continue
                img_bytes, img_ext = image_cache[img_hash]

                cached = self.cache.get(img_bytes)
                if cached:
                    description_cache[img_hash] = cached
                    cache_hits += 1
                    logger.debug("Cache hit for hash %s — skipping vision call.", img_hash[:8])
                    continue

                mime = "image/jpeg" if img_ext == "jpg" else f"image/{img_ext}"
                b64 = base64.b64encode(img_bytes).decode()
                data_uri = f"data:{mime};base64,{b64}"
                desc_tasks.append(
                    (img_hash, img_bytes, self._generate_image_description_async(data_uri))
                )

            logger.info(
                "Batch %d/%d | vision calls needed=%d | cache_hits=%d",
                batch_num, total_batches, len(desc_tasks), cache_hits,
            )

            described_count = 0
            fallback_count = 0
            results = await asyncio.gather(
                *(coro for _, _, coro in desc_tasks),
                return_exceptions=True,
            ) if desc_tasks else []

            for (img_hash, img_bytes, _), result in zip(desc_tasks, results):
                if isinstance(result, Exception):
                    logger.error(
                        "✗ Description failed for hash %s: %s", img_hash[:8], result
                    )
                    description_cache[img_hash] = self._fallback_response()
                    fallback_count += 1
                    continue

                if result:
                    description_cache[img_hash] = result
                    self.cache.set(img_bytes, result)
                    described_count += 1
                    logger.debug(
                        "Described hash %s | content_type=%s | is_action=%s",
                        img_hash[:8], result.content_type, result.is_action_step,
                    )
                else:
                    description_cache[img_hash] = self._fallback_response()
                    fallback_count += 1
                    logger.warning(
                        "No description returned for hash %s — using fallback.", img_hash[:8]
                    )

            logger.info(
                "Batch %d/%d complete | described=%d | fallbacks=%d | elapsed=%.2fs",
                batch_num, total_batches, described_count, fallback_count,
                time.monotonic() - t_batch,
            )

        logger.info(
            "Parallel processing done | total_extracted=%d | total_described=%d",
            len(image_cache), len(description_cache),
        )

    def _extract_image_from_pdf(self, img_dict: dict) -> Tuple[Optional[bytes], str]:
        """Extract raw image bytes from a pdfplumber image dict."""
        try:
            stream = img_dict.get("stream")
            if stream is None:
                logger.debug("No stream found in image dict.")
                return None, "png"

            raw = stream.get_data()
            filters = stream.get("/Filter", [])
            if isinstance(filters, str):
                filters = [filters]

            logger.debug("Image filters: %s | raw_size=%d bytes", filters, len(raw))

            # JPEG — send as-is
            if "/DCTDecode" in filters:
                logger.debug("JPEG detected — returning raw bytes.")
                return raw, "jpg"

            # JPEG 2000 — convert to PNG
            if "/JPXDecode" in filters:
                logger.debug("JPEG2000 detected — converting to PNG.")
                try:
                    buf = io.BytesIO()
                    Image.open(io.BytesIO(raw)).save(buf, format="PNG")
                    converted = buf.getvalue()
                    logger.debug("JPEG2000 → PNG | size=%d bytes", len(converted))
                    return converted, "png"
                except Exception as exc:
                    logger.error("JPEG2000 decode failed: %s", exc)
                    return None, "png"

            # Reconstruct from raw pixel data
            sw = stream.get("/Width")
            sh = stream.get("/Height")
            srcsize = img_dict.get("srcsize")

            if sw and sh:
                width, height = int(sw), int(sh)
                logger.debug("Image dimensions from stream: %dx%d", width, height)
            elif srcsize and len(srcsize) >= 2:
                width, height = int(srcsize[0]), int(srcsize[1])
                logger.debug("Image dimensions from srcsize: %dx%d", width, height)
            else:
                logger.debug("No explicit dimensions — attempting PIL open.")
                try:
                    buf = io.BytesIO()
                    Image.open(io.BytesIO(raw)).save(buf, format="PNG")
                    return buf.getvalue(), "png"
                except Exception:
                    logger.warning("PIL open fallback failed — no dimensions available.")
                    return None, "png"

            colorspace = str(img_dict.get("colorspace", ""))
            cs = stream.get("/ColorSpace")
            if cs:
                colorspace = str(cs)

            if "RGB" in colorspace or "DeviceRGB" in colorspace:
                mode, channels = "RGB", 3
            elif "CMYK" in colorspace or "DeviceCMYK" in colorspace:
                mode, channels = "CMYK", 4
            else:
                mode, channels = "L", 1

            logger.debug(
                "Colorspace=%s | mode=%s | channels=%d | expected_bytes=%d",
                colorspace, mode, channels, width * height * channels,
            )

            expected = width * height * channels
            if len(raw) >= expected:
                try:
                    pil = Image.frombytes(mode, (width, height), raw[:expected])
                    if mode == "CMYK":
                        pil = pil.convert("RGB")
                        logger.debug("Converted CMYK → RGB.")
                    buf = io.BytesIO()
                    pil.save(buf, format="PNG")
                    result = buf.getvalue()
                    logger.debug("Reconstructed image | png_size=%d bytes", len(result))
                    return result, "png"
                except Exception as exc:
                    logger.debug("frombytes reconstruction failed: %s", exc)

            # Absolute fallback
            logger.debug("Attempting absolute PIL fallback...")
            try:
                buf = io.BytesIO()
                Image.open(io.BytesIO(raw)).save(buf, format="PNG")
                return buf.getvalue(), "png"
            except Exception as exc:
                logger.error("PIL absolute fallback failed: %s", exc)
                return None, "png"

        except Exception as exc:
            logger.error("Error extracting PDF image: %s", exc)
            return None, "png"

    def _build_page_content_with_images(
        self,
        page: pdfplumber.page.Page,
        page_index: int,
        image_map: Dict,
        image_cache: Dict[str, Tuple[bytes, str]],
        description_cache: Dict[str, ImageDescriptionResponse],
    ) -> str:
        """
        Merge text words and image descriptions for a single PDF page,
        sorted by vertical position.
        """
        elements: List[Tuple[float, float, str, bool]] = []
        page_height = page.height

        # ── Text words → lines ───────────────────────────────────────────────
        words = page.extract_words(x_tolerance=3, y_tolerance=3) or []
        logger.debug("Page %d: extracted %d words.", page_index + 1, len(words))

        lines_map: Dict[int, List[dict]] = {}
        y_tol = 5
        for w in words:
            y0 = w["top"]
            key = next((k for k in lines_map if abs(k - y0) < y_tol), None)
            if key is not None:
                lines_map[key].append(w)
            else:
                lines_map[int(y0)] = [w]

        for key, line_words in lines_map.items():
            line_words.sort(key=lambda w: w["x0"])
            if not line_words:
                continue
            y = sum(w["top"] for w in line_words) / len(line_words)
            x = min(w["x0"] for w in line_words)
            text = " ".join(w["text"] for w in line_words)
            elements.append((y, x, text, False))

        logger.debug("Page %d: formed %d text lines.", page_index + 1, len(elements))

        # ── Images ───────────────────────────────────────────────────────────
        imgs_on_page = page.images or []
        inserted_img_count = 0
        skipped_img_count = 0

        for img in imgs_on_page:
            x0, y0, x1, y1 = img["x0"], img["top"], img["x1"], img["bottom"]
            if (x1 - x0) < self.config.min_image_size and (y1 - y0) < self.config.min_image_size:
                skipped_img_count += 1
                continue

            stream = img.get("stream")
            if stream is None:
                skipped_img_count += 1
                continue

            try:
                raw = stream.get_data()
                img_hash = hashlib.md5(raw).hexdigest()
            except Exception:
                skipped_img_count += 1
                continue

            if img_hash not in image_cache or img_hash not in description_cache:
                logger.debug(
                    "Page %d: hash %s not in cache — skipping.", page_index + 1, img_hash[:8]
                )
                skipped_img_count += 1
                continue

            desc = description_cache[img_hash]
            if not desc.has_content or desc.is_decorative:
                logger.debug(
                    "Page %d: Image hash %s is decorative or has no content — skipping.",
                    page_index + 1, img_hash[:8],
                )
                skipped_img_count += 1
                continue

            if self.config.filter_headers_footers and page_height:
                y_ratio = ((y0 + y1) / 2) / page_height
                if y_ratio < 0.08 or y_ratio > 0.92:
                    logger.debug(
                        "Page %d: Image at y_ratio=%.2f filtered as header/footer.",
                        page_index + 1, y_ratio,
                    )
                    skipped_img_count += 1
                    continue

            desc_text = self._format_image_description(desc)
            if desc_text.strip():
                elements.append(((y0 + y1) / 2, x0, desc_text, True))
                inserted_img_count += 1

        logger.debug(
            "Page %d: images inserted=%d | skipped=%d",
            page_index + 1, inserted_img_count, skipped_img_count,
        )

        # ── Sort & merge ─────────────────────────────────────────────────────
        elements.sort(key=lambda e: (e[0], e[1]))
        if not elements:
            logger.debug("Page %d: no content elements after filtering.", page_index + 1)
            return ""

        text_ys = sorted(e[0] for e in elements if not e[3])
        if len(text_ys) >= 2:
            gaps = [text_ys[i + 1] - text_ys[i] for i in range(len(text_ys) - 1)]
            typical_lh = sum(gaps) / len(gaps)
        else:
            typical_lh = 15
        para_threshold = typical_lh * 1.5

        logger.debug(
            "Page %d: typical_line_height=%.1f | para_threshold=%.1f",
            page_index + 1, typical_lh, para_threshold,
        )

        parts = []
        prev_y = None
        for y, x, content, is_image in elements:
            content = content.strip()
            if not content:
                continue
            if prev_y is None:
                parts.append(content)
            elif (y - prev_y) > para_threshold:
                parts.append("\n\n" + content)
            else:
                parts.append(" " + content)
            prev_y = y

        result = "".join(parts)
        logger.debug("Page %d: final content length=%d chars.", page_index + 1, len(result))
        return result

    async def _extract_pdf_text_only(self, pdf_path: str) -> str:
        """Fallback: extract plain text from every PDF page."""
        logger.info("Text-only PDF extraction: %s", pdf_path)
        t_start = time.monotonic()
        pdf = None
        try:
            pdf = pdfplumber.open(pdf_path)
            parts = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    parts.append(page_text)
            pdf.close()
            pdf = None
            result = "\n\n".join(parts)
            logger.info(
                "Text-only extraction complete | pages_with_text=%d | chars=%d | elapsed=%.2fs",
                len(parts), len(result), time.monotonic() - t_start,
            )
            return result
        except Exception as exc:
            logger.error("Text-only PDF extraction failed: %s | elapsed=%.2fs", exc, time.monotonic() - t_start)
            return ""
        finally:
            if pdf:
                try:
                    pdf.close()
                except Exception:
                    pass

    # ──────────────────────────────────────────────────────────────────────────
    # DOCX internals
    # ──────────────────────────────────────────────────────────────────────────

    def _extract_docx_elements_sync(self, docx_path: str) -> List[TextElement]:
        """
        Walk a DOCX document body and return a flat list of TextElement objects
        (text paragraphs and embedded images) with sequential position values.
        """
        if not DOCX_AVAILABLE:
            return []

        logger.info("Extracting DOCX elements (sync) from: %s", docx_path)
        elements: List[TextElement] = []
        position = 0.0

        text_count = 0
        image_count = 0
        skipped_small_images = 0
        table_count = 0

        try:
            doc = docx.Document(docx_path)

            # Build relationship map: rId -> image part
            image_rels: Dict[str, Any] = {
                rel.rId: rel.target_part
                for rel in doc.part.rels.values()
                if "image" in rel.reltype
            }
            logger.info("DOCX: found %d image relationships.", len(image_rels))

            NS_WD = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            NS_A  = "http://schemas.openxmlformats.org/drawingml/2006/main"
            NS_R  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

            for element in doc.element.body:
                tag = element.tag

                # ── Paragraph ────────────────────────────────────────────────
                if tag.endswith("p"):
                    drawings = (
                        element.findall(f".//{{{NS_WP}}}inline")
                        + element.findall(f".//{{{NS_WP}}}anchor")
                    )
                    for drawing in drawings:
                        blip = drawing.find(f".//{{{NS_A}}}blip")
                        if blip is None:
                            continue
                        embed_id = blip.get(f"{{{NS_R}}}embed")
                        if not embed_id or embed_id not in image_rels:
                            logger.debug("Drawing with embed_id=%s not in image_rels — skipping.", embed_id)
                            continue

                        image_part = image_rels[embed_id]
                        image_bytes = image_part.blob
                        if len(image_bytes) < 100:
                            logger.debug("Skipping tiny image blob (%d bytes).", len(image_bytes))
                            skipped_small_images += 1
                            continue

                        alt_text = ""
                        doc_pr = drawing.find(f".//{{{NS_WP}}}docPr")
                        if doc_pr is not None:
                            alt_text = doc_pr.get("descr", "") or doc_pr.get("title", "")

                        content_type = getattr(image_part, "content_type", "image/png")
                        logger.debug(
                            "Found image | embed_id=%s | size=%d bytes | mime=%s | alt=%r",
                            embed_id, len(image_bytes), content_type, alt_text,
                        )

                        elements.append(TextElement(
                            y=position,
                            x=0,
                            content=base64.b64encode(image_bytes).decode(),
                            element_type="image",
                            alt_text=alt_text,
                            content_type_mime=content_type,
                        ))
                        image_count += 1

                    # Paragraph text
                    para_text = "".join(element.itertext()).strip()
                    if para_text:
                        elements.append(
                            TextElement(y=position, x=0, content=para_text, element_type="text")
                        )
                        text_count += 1
                    position += 1.0

                # ── Table ─────────────────────────────────────────────────────
                elif tag.endswith("tbl"):
                    table_text = self._extract_table_text(element)
                    if table_text:
                        elements.append(
                            TextElement(y=position, x=0, content=table_text, element_type="text")
                        )
                        table_count += 1
                    position += 1.0

        except Exception as exc:
            logger.error("Error extracting DOCX elements: %s", exc, exc_info=True)

        logger.info(
            "DOCX element extraction done | text_paras=%d | tables=%d | images=%d | skipped_images=%d",
            text_count, table_count, image_count, skipped_small_images,
        )
        return elements

    def _extract_table_text(self, table_element) -> str:
        """Convert a DOCX table XML element to a pipe-delimited string."""
        NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        rows = []
        for row in table_element.findall(f".//{{{NS}}}tr"):
            cells = [
                "".join(cell.itertext()).strip()
                for cell in row.findall(f".//{{{NS}}}tc")
            ]
            if any(cells):
                rows.append(" | ".join(cells))
        logger.debug("Extracted table: %d rows.", len(rows))
        return "\n".join(rows)

    async def _process_docx_images(
        self, image_elements: List[TextElement]
    ) -> Dict[int, str]:
        """Generate descriptions for all DOCX image elements."""
        descriptions: Dict[int, str] = {}
        batch_size = self.config.max_concurrent_images
        total_batches = (len(image_elements) + batch_size - 1) // batch_size

        logger.info(
            "Processing DOCX images | total=%d | batch_size=%d | batches=%d",
            len(image_elements), batch_size, total_batches,
        )

        for batch_start in range(0, len(image_elements), batch_size):
            batch = image_elements[batch_start : batch_start + batch_size]
            batch_num = batch_start // batch_size + 1
            logger.info(
                "DOCX image batch %d/%d | images=%d",
                batch_num, total_batches, len(batch),
            )
            t_batch = time.monotonic()

            coros = [
                self._generate_docx_image_description(
                    batch_start + i,
                    f"data:{elem.content_type_mime};base64,{elem.content}",
                    elem.alt_text,
                )
                for i, elem in enumerate(batch)
            ]
            results = await asyncio.gather(*coros, return_exceptions=True)

            success = 0
            failed = 0
            for i, result in enumerate(results):
                global_idx = batch_start + i
                if isinstance(result, Exception):
                    failed += 1
                    logger.error("DOCX image %d failed: %s", global_idx, result)
                    alt = image_elements[global_idx].alt_text
                    fb = self._format_fallback_description(alt)
                    if fb:
                        descriptions[global_idx] = fb
                elif result:
                    descriptions[global_idx] = result
                    success += 1
                    logger.debug("DOCX image %d described successfully.", global_idx)

            logger.info(
                "DOCX batch %d/%d done | success=%d | failed=%d | elapsed=%.2fs",
                batch_num, total_batches, success, failed, time.monotonic() - t_batch,
            )

        logger.info(
            "DOCX image processing complete | total_described=%d", len(descriptions)
        )
        return descriptions

    async def _generate_docx_image_description(
        self, idx: int, data_uri: str, alt_text: str
    ) -> Optional[str]:
        """Describe a single DOCX image, using cache where possible."""
        try:
            image_bytes = base64.b64decode(data_uri.split(",", 1)[1])
            logger.debug("DOCX image %d | size=%d bytes | alt=%r", idx, len(image_bytes), alt_text)

            cached = self.cache.get(image_bytes)
            if cached:
                logger.debug("DOCX image %d: cache hit.", idx)
                return self._format_image_description(cached, alt_text) or None

            logger.debug("DOCX image %d: calling vision model...", idx)
            t_start = time.monotonic()
            response = await self._generate_image_description_async(
                data_uri, f"Image from document (alt: {alt_text})"
            )
            elapsed = time.monotonic() - t_start

            if response:
                self.cache.set(image_bytes, response)
                logger.debug(
                    "DOCX image %d: described | content_type=%s | elapsed=%.2fs",
                    idx, response.content_type, elapsed,
                )
                return self._format_image_description(response, alt_text) or None

            logger.warning("DOCX image %d: no response from vision model — using fallback.", idx)
            return self._format_fallback_description(alt_text) if alt_text else None

        except Exception as exc:
            logger.error("Error describing DOCX image %d: %s", idx, exc)
            return self._format_fallback_description(alt_text) if alt_text else None

    def _merge_docx_elements(
        self,
        text_elements: List[TextElement],
        image_descriptions: Dict[int, str],
    ) -> str:
        """
        Interleave text paragraphs and image descriptions sorted by position.
        """
        logger.info(
            "Merging DOCX elements | text=%d | image_descriptions=%d",
            len(text_elements), len(image_descriptions),
        )

        all_elements: List[Tuple[float, bool, str]] = []

        for elem in text_elements:
            content = (elem.content or "").strip()
            if content:
                all_elements.append((elem.y, False, content))

        for idx, desc in image_descriptions.items():
            desc_clean = (desc or "").strip()
            if desc_clean:
                all_elements.append((idx * 0.5, True, desc_clean))

        all_elements.sort(key=lambda x: x[0])
        if not all_elements:
            logger.warning("No elements to merge — returning empty string.")
            return ""

        text_ys = sorted(e[0] for e in all_elements if not e[1])
        if len(text_ys) >= 2:
            gaps = [text_ys[i + 1] - text_ys[i] for i in range(len(text_ys) - 1)]
            typical_gap = sum(gaps) / len(gaps)
        else:
            typical_gap = 1.0
        para_threshold = typical_gap * 1.5

        logger.debug(
            "Merge params | typical_gap=%.2f | para_threshold=%.2f",
            typical_gap, para_threshold,
        )

        parts = []
        prev_y = None
        for y, is_image, content in all_elements:
            if not content:
                continue
            if prev_y is None:
                parts.append(content)
            elif (y - prev_y) > para_threshold:
                parts.append("\n\n" + content)
            else:
                parts.append(" " + content)
            prev_y = y

        result = "".join(parts)
        logger.info("Merge complete | output_length=%d chars.", len(result))
        return result

    async def _extract_docx_text_only(self, docx_path: str) -> str:
        """Fallback: extract plain text from DOCX without image processing."""
        if not DOCX_AVAILABLE:
            return ""
        logger.info("Text-only DOCX extraction: %s", docx_path)
        try:
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(
                self.thread_pool, self._extract_docx_text_sync, docx_path
            )
        except Exception as exc:
            logger.error("Text-only DOCX extraction failed: %s", exc)
            return ""

    def _extract_docx_text_sync(self, docx_path: str) -> str:
        try:
            doc = docx.Document(docx_path)
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            result = "\n\n".join(paras)
            logger.info("Text-only DOCX sync extraction | paragraphs=%d | chars=%d", len(paras), len(result))
            return result
        except Exception as exc:
            logger.error("Sync DOCX text extraction failed: %s", exc)
            return ""

    # ──────────────────────────────────────────────────────────────────────────
    # HTML internals
    # ──────────────────────────────────────────────────────────────────────────

    async def _get_image_data(self, src: str, base_url: Optional[str]) -> Optional[bytes]:
        """Resolve and download image bytes from a src attribute."""
        try:
            if src.startswith("data:"):
                logger.debug("Decoding inline data URI (%d chars).", len(src))
                return self._decode_data_uri(src)
            elif src.startswith(("http://", "https://")):
                logger.debug("Downloading remote image: %s", src[:80])
                return await self._download_image(src)
            elif base_url:
                resolved = urljoin(base_url, src)
                logger.debug("Resolved relative URL: %s → %s", src[:60], resolved[:80])
                return await self._download_image(resolved)
            else:
                logger.warning("Cannot resolve image source (no base_url): %s", src[:80])
                return None
        except Exception as exc:
            logger.warning("Error getting image data from %s: %s", src[:80], exc)
            return None

    def _decode_data_uri(self, data_uri: str) -> Optional[bytes]:
        try:
            _, data = data_uri.split(",", 1)
            result = base64.b64decode(data)
            logger.debug("Decoded data URI | size=%d bytes.", len(result))
            return result
        except Exception as exc:
            logger.warning("Error decoding data URI: %s", exc)
            return None

    async def _download_image(self, url: str) -> Optional[bytes]:
        logger.debug("Downloading image: %s", url[:80])
        t_start = time.monotonic()
        retryable_statuses = {408, 425, 429, 500, 502, 503, 504}

        for attempt in range(1, self.config.download_max_retries + 1):
            try:
                async with aiohttp.ClientSession() as session:
                    async with session.get(
                        url,
                        timeout=aiohttp.ClientTimeout(total=self.config.request_timeout),
                    ) as resp:
                        if resp.status == 200:
                            data = await resp.read()
                            logger.debug(
                                "Downloaded %d bytes from %s | elapsed=%.2fs | attempt=%d",
                                len(data), url[:60], time.monotonic() - t_start, attempt,
                            )
                            return data

                        should_retry = (
                            resp.status in retryable_statuses
                            and attempt < self.config.download_max_retries
                        )
                        logger.warning(
                            "HTTP %d for image URL: %s | attempt=%d/%d | retry=%s | elapsed=%.2fs",
                            resp.status,
                            url[:80],
                            attempt,
                            self.config.download_max_retries,
                            should_retry,
                            time.monotonic() - t_start,
                        )
                        if not should_retry:
                            return None
            except Exception as exc:
                should_retry = attempt < self.config.download_max_retries
                logger.warning(
                    "Error downloading image from %s: %s | attempt=%d/%d | retry=%s | elapsed=%.2fs",
                    url[:80],
                    exc,
                    attempt,
                    self.config.download_max_retries,
                    should_retry,
                    time.monotonic() - t_start,
                )
                if not should_retry:
                    return None

            await asyncio.sleep(self._retry_delay(attempt))

        return None

    def _guess_mime_type(self, src: str) -> str:
        if src.startswith("data:"):
            try:
                return src.split(",")[0].split(":")[1].split(";")[0]
            except Exception:
                return "image/png"
        lower = src.lower()
        if ".jpg" in lower or ".jpeg" in lower:
            return "image/jpeg"
        if ".gif" in lower:
            return "image/gif"
        if ".webp" in lower:
            return "image/webp"
        return "image/png"

    def _clean_html_with_descriptions(
        self, soup: BeautifulSoup, image_descriptions: Dict[int, str]
    ) -> str:
        """Replace <img> tags with descriptions and clean up the HTML tree."""
        result_parts: List[str] = []
        img_index = 0

        BLOCK = {
            "p", "div", "h1", "h2", "h3", "h4", "h5", "h6",
            "article", "section", "ol", "ul",
        }
        SKIP = {"script", "style"}

        def process(element):
            nonlocal img_index

            if not hasattr(element, "name") or element.name is None:
                text = str(element).strip()
                if text:
                    result_parts.append(text + " ")
                return

            name = element.name
            if name == "img":
                if img_index in image_descriptions:
                    result_parts.append(image_descriptions[img_index])
                    logger.debug("Inserted description for image index %d.", img_index)
                else:
                    logger.debug("No description for image index %d — omitted.", img_index)
                img_index += 1
            elif name == "br":
                result_parts.append("\n")
            elif name in SKIP:
                pass
            elif name == "li":
                result_parts.append("\n* ")
                for child in element.children:
                    process(child)
            elif name in BLOCK:
                result_parts.append("\n")
                for child in element.children:
                    process(child)
                result_parts.append("\n")
            else:
                for child in element.children:
                    process(child)

        body = soup.find("body") or soup
        process(body)

        text = "".join(result_parts)
        lines = [line.strip() for line in text.splitlines()]
        result = "\n".join(line for line in lines if line)
        logger.debug(
            "HTML tree walk complete | img_tags_encountered=%d | output_lines=%d",
            img_index, len(lines),
        )
        return result

    def _clean_html_text_only(self, html_content: str) -> str:
        if not html_content:
            return ""
        logger.info("Cleaning HTML text only (no image processing) | input_length=%d", len(html_content))
        try:
            soup = BeautifulSoup(html_content, "html.parser")
            for el in soup(["script", "style"]):
                el.decompose()
            lines = [line.strip() for line in soup.get_text().splitlines()]
            result = "\n".join(line for line in lines if line)
            logger.info("HTML text-only clean | output_length=%d chars.", len(result))
            return result
        except Exception as exc:
            logger.warning("Error cleaning HTML: %s", exc)
            return html_content

    # ──────────────────────────────────────────────────────────────────────────
    # Vision API (Azure OpenAI)
    # ──────────────────────────────────────────────────────────────────────────

    async def _generate_image_description_async(
        self,
        image_data_uri: str,
        context: Optional[str] = None,
    ) -> Optional[ImageDescriptionResponse]:
        """Send image to Azure OpenAI vision model and return structured description."""
        if not self._client:
            logger.warning("Vision client not available — skipping description.")
            return None

        try:
            prompt = self._build_image_prompt(context)
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {"url": image_data_uri, "detail": "high"},
                        },
                    ],
                }
            ]

            logger.debug(
                "Submitting vision call | context=%r | uri_length=%d",
                context[:60] if context else None, len(image_data_uri),
            )
            t_start = time.monotonic()
            loop = asyncio.get_running_loop()
            response = await loop.run_in_executor(
                self.thread_pool, self._call_vision_model_sync, messages
            )
            elapsed = time.monotonic() - t_start

            if response:
                response.text = re.sub(r"\s+", " ", response.text.strip())
                response.text = (
                    response.text
                    .replace("[", "(").replace("]", ")")
                    .replace("`", "'").replace('"', "'")
                    .replace("\\", "")
                )
                logger.debug(
                    "Vision call ✓ | content_type=%s | is_action=%s | shows_result=%s | "
                    "has_instructions=%s | elapsed=%.2fs",
                    response.content_type, response.is_action_step,
                    response.shows_result, response.contains_instructions, elapsed,
                )
            else:
                logger.warning("Vision call returned None | elapsed=%.2fs", elapsed)

            return response

        except Exception as exc:
            logger.error("Error generating image description: %s", exc, exc_info=True)
            return None

    def _call_vision_model_sync(
        self, messages: List[Dict[str, Any]]
    ) -> Optional[ImageDescriptionResponse]:
        """
        Synchronous Azure OpenAI vision model call (runs in thread pool).
        """
        for attempt in range(1, self.config.vision_max_retries + 1):
            try:
                logger.debug(
                    "Vision model sync call | deployment=%s | temperature=%.2f | attempt=%d/%d",
                    env_config.LLM_AZURE_OPENAI_DEPLOYMENT_NAME,
                    self.config.temperature,
                    attempt,
                    self.config.vision_max_retries,
                )
                result = self._client.chat.completions.create(
                    model=env_config.LLM_AZURE_OPENAI_DEPLOYMENT_NAME,
                    messages=messages,
                    temperature=self.config.temperature,
                    response_model=ImageDescriptionResponse,
                )
                logger.debug("Vision model sync call returned successfully.")
                return result
            except Exception as exc:
                should_retry = attempt < self.config.vision_max_retries
                logger.warning(
                    "Vision model call failed | attempt=%d/%d | retry=%s | error=%s",
                    attempt,
                    self.config.vision_max_retries,
                    should_retry,
                    exc,
                )
                if not should_retry:
                    return None
                time.sleep(self._retry_delay(attempt))

        return None

    def _retry_delay(self, attempt: int) -> float:
        return self.config.retry_base_delay_sec * (2 ** max(0, attempt - 1))

    def _build_image_prompt(self, context: Optional[str] = None) -> str:
        prompt = """Analyze this image from a documentation or knowledge base article.

YOUR TASK: Determine if this image shows an ACTION the user needs to take, or if it shows a RESULT/CONFIRMATION.

DECISION RULES:
- Choose exactly ONE primary content_type for the image.
- content_type must be one of: text_extraction, action_step, dialog_box, diagram, screenshot, result_display, instructions, unknown.
- If more than one category seems possible, use this precedence:
  instructions > dialog_box > action_step > result_display > diagram > screenshot > text_extraction > unknown
- Set is_action_step=True only when the image itself shows something the user still needs to do now.
- Set shows_result=True only when the image primarily shows an outcome, status, error, or confirmation rather than an action still to be taken.
- A dialog can still be an action step if the user must click/select something, but the primary content_type should remain dialog_box.
- If the image is mostly text instructions, the primary content_type must be instructions even if buttons or UI are visible.

CLASSIFICATION:
1. ACTION STEP (is_action_step=True): Image shows something the user must DO
   - A button they need to click (e.g., "Click Accept", "Click Sign In")
   - A field they need to fill (e.g., "Enter your username")
   - A checkbox they need to check/uncheck
   - A dropdown or option they need to select
   - Set action_summary to a concise instruction like "Click the Accept button"

2. RESULT/CONFIRMATION (shows_result=True): Image shows an OUTCOME
   - Success/confirmation messages ("Connected successfully")
   - Error messages or warnings
   - Status indicators
   - Completed state after an action

3. DIALOG BOX (content_type="dialog_box"): Popup or modal with options
   - Identify the dialog title and message
   - List all available buttons/options in ui_elements
   - If user needs to click something, set is_action_step=True

4. INSTRUCTIONS/TEXT CONTENT (contains_instructions=True): Image contains written steps
   - Screenshot of a document, email, or text-based instructions
   - Numbered or bulleted list of steps shown in the image
   - Error messages, code snippets, or configuration text
   - Set content_type="instructions" and extract ALL text into extracted_text verbatim

5. INFORMATIONAL: General screenshot, diagram, or reference image

UI ELEMENTS (ui_elements array): capture button labels, checkbox labels, link text, input field labels.

TEXT EXTRACTION: If image contains readable steps or significant text, set contains_instructions=True
and put COMPLETE verbatim text in extracted_text, preserving numbering and structure.

EMPTY/DECORATIVE: has_content=False for blank images; is_decorative=True for purely ornamental images.

Return values that are internally consistent:
- If has_content=False, then is_action_step=False and shows_result=False.
- If is_decorative=True, then has_content should usually be False.
- If contains_instructions=True, extracted_text should not be empty.
- If is_action_step=True, action_summary should be a short imperative instruction.

Use plain text only — no brackets, backticks, or special formatting."""

        if context:
            prompt += f"\n\nSurrounding document context:\n{context[:500]}"

        return prompt

    # ──────────────────────────────────────────────────────────────────────────
    # Formatting helpers
    # ──────────────────────────────────────────────────────────────────────────

    def _format_image_description(
        self,
        response: ImageDescriptionResponse,
        fallback_alt: str = "",
    ) -> str:
        """Convert a structured ImageDescriptionResponse to a plain-text string."""
        if not response.has_content or response.is_decorative:
            logger.debug("Image is decorative or has no content — returning empty string.")
            return ""

        if response.contains_instructions and response.extracted_text:
            logger.debug("Returning extracted instructions text.")
            return response.extracted_text.strip()

        if response.is_action_step and response.action_summary:
            parts = [f"[Action: {response.action_summary}]"]
            if response.text and response.text.lower() != response.action_summary.lower():
                parts.append(response.text)
            result = " ".join(parts)
            logger.debug("Formatted as action step: %s", result[:80])
            return result

        if response.shows_result:
            result = f"[Result: {response.text}]"
            logger.debug("Formatted as result: %s", result[:80])
            return result

        if response.content_type == "dialog_box":
            desc = response.text
            if response.ui_elements:
                desc += f" Options available: {', '.join(response.ui_elements[:5])}"
            result = f"[Dialog: {desc}]"
            logger.debug("Formatted as dialog: %s", result[:80])
            return result

        formatted = f"[Image: {response.text}]"
        if response.ui_elements:
            formatted += f" (UI elements: {', '.join(response.ui_elements[:5])})"
        logger.debug("Formatted as general image: %s", formatted[:80])
        return formatted

    def _format_fallback_description(self, alt_text: str = "") -> str:
        result = f"[Image: {alt_text}]" if alt_text else "[Image: Visual content]"
        logger.debug("Fallback description: %s", result)
        return result

    def _fallback_response(self) -> ImageDescriptionResponse:
        logger.debug("Creating fallback ImageDescriptionResponse.")
        return ImageDescriptionResponse(
            text="Visual content (image)",
            has_content=True,
            is_decorative=False,
            content_type="unknown",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# Convenience async functions
# ═══════════════════════════════════════════════════════════════════════════════

async def process_pdf_with_images(
    pdf_path: str,
    config: Optional[ImageProcessingConfig] = None,
) -> str:
    """Process a PDF file with image extraction and description."""
    processor = ImageProcessor(config=config)
    try:
        return await processor.process_pdf_with_images(pdf_path)
    finally:
        processor.shutdown()


async def process_docx_with_images(
    docx_path: str,
    config: Optional[ImageProcessingConfig] = None,
) -> str:
    """Process a DOCX file with image extraction and description."""
    processor = ImageProcessor(config=config)
    try:
        return await processor.process_docx_with_images(docx_path)
    finally:
        processor.shutdown()


async def process_html_with_images(
    html_content: str,
    base_url: Optional[str] = None,
    config: Optional[ImageProcessingConfig] = None,
) -> str:
    """Process HTML content with image extraction and description."""
    processor = ImageProcessor(config=config)
    try:
        return await processor.process_html_with_images(html_content, base_url)
    finally:
        processor.shutdown()


# ═══════════════════════════════════════════════════════════════════════════════
# CLI entry point
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import sys

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s | %(message)s",
    )

    async def main():
        cfg = ImageProcessingConfig(
            min_image_size=32,
            max_images_per_doc=30,
            filter_headers_footers=True,
            filter_decorative=True,
        )

        processor = ImageProcessor(config=cfg)

        try:
            if len(sys.argv) < 2:
                print("Usage: python image_processor.py <file.pdf|file.docx|file.html>")
                return

            path = sys.argv[1]
            ext = Path(path).suffix.lower()

            logger.info("CLI invocation | file=%s | type=%s", path, ext)

            if ext == ".pdf":
                result = await processor.process_pdf_with_images(path)
            elif ext in (".docx", ".doc"):
                result = await processor.process_docx_with_images(path)
            elif ext in (".html", ".htm"):
                html = Path(path).read_text(encoding="utf-8")
                result = await processor.process_html_with_images(
                    html,
                    base_url=f"file://{Path(path).resolve().parent}/",
                )
            else:
                logger.error("Unsupported file type: %s", ext)
                print(f"Unsupported file type: {ext}")
                return

            logger.info("Output length: %d chars", len(result))
            print(result)

        finally:
            processor.shutdown()

    asyncio.run(main())
